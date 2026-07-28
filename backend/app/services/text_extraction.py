import os
from typing import Optional
import pdfplumber
import docx
import pytesseract
from PIL import Image

def extract_text_from_file(filepath: str, file_type: str) -> Optional[str]:
    """
    Routes the file to the appropriate text extraction method based on mime type or extension.
    """
    if not os.path.exists(filepath):
        return None

    file_ext = os.path.splitext(filepath)[1].lower()
    
    try:
        if file_type == 'application/pdf' or file_ext == '.pdf':
            return _extract_from_pdf(filepath)
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_ext == '.docx':
            return _extract_from_docx(filepath)
        elif file_type.startswith('image/') or file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            return _extract_from_image(filepath)
        elif file_type.startswith('text/') or file_ext in ['.txt', '.csv']:
            return _extract_from_text(filepath)
        else:
            return None
    except Exception as e:
        print(f"Error extracting text from {filepath}: {e}")
        return None

def _extract_from_pdf(filepath: str) -> str:
    extracted_text = []
    
    # 1. Try PyMuPDF (fitz) first
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        for page in doc:
            t = page.get_text()
            if t and t.strip():
                extracted_text.append(t.strip())
        doc.close()
    except Exception as e:
        print(f"PyMuPDF extraction failed: {e}")

    # 2. Fallback to pdfplumber if fitz produced no text
    if not extracted_text:
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        extracted_text.append(text.strip())
        except Exception as e:
            print(f"pdfplumber extraction failed: {e}")

    full_text = "\n".join(extracted_text).strip()

    # 3. If scanned PDF (no text found), run pytesseract OCR on page images
    if not full_text:
        try:
            import fitz
            doc = fitz.open(filepath)
            ocr_text = []
            for page in doc:
                pix = page.get_pixmap(dpi=150)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                txt = pytesseract.image_to_string(img)
                if txt and txt.strip():
                    ocr_text.append(txt.strip())
            doc.close()
            full_text = "\n".join(ocr_text).strip()
        except Exception as e:
            print(f"Scanned PDF OCR failed: {e}")

    return full_text

def _extract_from_docx(filepath: str) -> str:
    extracted_text = []
    doc = docx.Document(filepath)
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            extracted_text.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                extracted_text.append(" | ".join(row_text))
    return "\n".join(extracted_text)

def _extract_from_image(filepath: str) -> str:
    try:
        image = Image.open(filepath)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        print(f"OCR Error extracting image text: {e}")
        return ""

def _extract_from_text(filepath: str) -> str:
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

